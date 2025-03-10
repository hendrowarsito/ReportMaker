import streamlit as st
from docx import Document
from io import BytesIO
import pandas as pd
import zipfile
import locale
import os

def format_number_indonesia(value):
    """Format number to Indonesian format (e.g., 12.000,00)."""
    try:
        locale.setlocale(locale.LC_NUMERIC, "id_ID.UTF-8")
        return locale.format_string("%.2f", value, grouping=True)
    except:
        return value

def replace_placeholders(document, replacements):
    """Replace placeholders in a DOCX document with provided values."""
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            formatted_value = format_number_indonesia(value) if isinstance(value, (int, float)) else value
            paragraph.text = paragraph.text.replace(f"{{{key}}}", str(formatted_value))
    return document

def extract_placeholders(document):
    """Extract placeholders from a DOCX document."""
    placeholders = set()
    for paragraph in document.paragraphs:
        if "{" in paragraph.text and "}" in paragraph.text:
            placeholders.update(
                part.strip("{}") for part in paragraph.text.split() if part.startswith("{") and part.endswith("}")
            )
    return sorted(placeholders)

def save_docx(document):
    """Save changes to a new DOCX file."""
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_zip(files):
    """Generate a ZIP file from a list of files."""
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_name, file_buffer in files:
            zf.writestr(file_name, file_buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer

def main():
    st.title("SRR KALIBATA REPORT MAKER")
    st.write("Upload DOCX templates and an Excel file to generate reports automatically.")

    uploaded_templates = st.file_uploader("Upload DOCX Templates", type="docx", accept_multiple_files=True)
    uploaded_excel = st.file_uploader("Upload Excel File", type="xlsx")
    
    if uploaded_templates and uploaded_excel:
        st.success(f"{len(uploaded_templates)} templates uploaded successfully!")
        data = pd.read_excel(uploaded_excel)
        st.write("Data Preview:")
        st.dataframe(data)

        templates = {}
        all_placeholders = set()
        for file in uploaded_templates:
            document = Document(file)
            placeholders = extract_placeholders(document)
            templates[file.name] = {"document": document, "placeholders": placeholders}
            all_placeholders.update(placeholders)

        unmatched_placeholders = [ph for ph in all_placeholders if ph not in data.columns]
        if unmatched_placeholders:
            st.warning(f"Unmatched placeholders: {', '.join(unmatched_placeholders)}")

        if st.button("Generate Reports"):
            st.success("Generating reports...")
            generated_files = []
            for index, row in data.iterrows():
                for template_name, template_data in templates.items():
                    document = replace_placeholders(template_data["document"], row.to_dict())
                    file_name = f"{index+1}_{template_name}"
                    buffer = save_docx(document)
                    generated_files.append((file_name, buffer))

            zip_buffer = generate_zip(generated_files)
            st.download_button("Download All Reports as ZIP", data=zip_buffer, file_name="generated_reports.zip", mime="application/zip")

if __name__ == "__main__":
    main()
